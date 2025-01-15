import sys
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QPushButton, QLabel, QFileDialog,
    QTextEdit, QVBoxLayout, QWidget, QHBoxLayout, QSlider
)
from PyQt5.QtGui import QFont
from PyQt5.QtCore import Qt
import fitz  # PyMuPDF
from transformers import pipeline
import re
import spacy
import importlib.util
from docx import Document
from PIL import Image
import pytesseract

class PDFSummarizerApp(QMainWindow):
    def __init__(self):
        super().__init__()

        # Check for required deep learning framework
        self.summarizer = None
        if self.check_framework_installed("torch"):
            try:
                self.summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6", framework="pt")
            except Exception as e:
                self.show_error(f"Failed to load summarisation model: {str(e)}")
        else:
            self.show_error("PyTorch is not installed. Summarisation features will be unavailable.")

        # Initialize spaCy for keyword extraction
        self.nlp = None
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            self.show_error("The spaCy model 'en_core_web_sm' is not installed. Please run 'python -m spacy download en_core_web_sm' to install it.")

        # Default maximum summary length
        self.max_summary_length = 100

        # Set up the main window
        self.setWindowTitle("PDF & Word Summarizer & Stats Extractor")
        self.setGeometry(100, 100, 800, 600)
        self.init_ui()

    def check_framework_installed(self, framework_name):
        return importlib.util.find_spec(framework_name) is not None

    def show_error(self, message):
        error_dialog = QLabel(message)
        error_dialog.setFont(QFont("Arial", 12))
        error_dialog.setAlignment(Qt.AlignCenter)
        error_dialog.setStyleSheet("color: red;")
        error_dialog.show()

    def init_ui(self):
        # Layouts
        main_layout = QVBoxLayout()
        button_layout = QHBoxLayout()
        slider_layout = QHBoxLayout()

        # Buttons
        self.open_button = QPushButton("Open File")
        self.open_button.setFont(QFont("Arial", 12))
        self.open_button.clicked.connect(self.open_file)

        self.summarize_button = QPushButton("Summarize")
        self.summarize_button.setFont(QFont("Arial", 12))
        self.summarize_button.clicked.connect(self.summarize_text)
        self.summarize_button.setEnabled(False)

        self.extract_stats_button = QPushButton("Extract Stats")
        self.extract_stats_button.setFont(QFont("Arial", 12))
        self.extract_stats_button.clicked.connect(self.extract_stats)
        self.extract_stats_button.setEnabled(False)

        button_layout.addWidget(self.open_button)
        button_layout.addWidget(self.summarize_button)
        button_layout.addWidget(self.extract_stats_button)

        # Summary length slider
        slider_label = QLabel("Summary Length:")
        slider_label.setFont(QFont("Arial", 12))

        self.length_slider = QSlider(Qt.Horizontal)
        self.length_slider.setMinimum(50)
        self.length_slider.setMaximum(300)
        self.length_slider.setValue(100)
        self.length_slider.setTickInterval(10)
        self.length_slider.setTickPosition(QSlider.TicksBelow)
        self.length_slider.valueChanged.connect(self.update_summary_length)

        self.slider_value_label = QLabel("100 words")
        self.slider_value_label.setFont(QFont("Arial", 12))

        slider_layout.addWidget(slider_label)
        slider_layout.addWidget(self.length_slider)
        slider_layout.addWidget(self.slider_value_label)

        # Text display areas
        self.text_display = QTextEdit()
        self.text_display.setFont(QFont("Arial", 11))
        self.text_display.setReadOnly(True)

        self.output_display = QTextEdit()
        self.output_display.setFont(QFont("Arial", 11))
        self.output_display.setReadOnly(True)

        # Labels
        input_label = QLabel("Extracted Text:")
        input_label.setFont(QFont("Arial", 14))
        input_label.setAlignment(Qt.AlignLeft)

        output_label = QLabel("Summary/Stats:")
        output_label.setFont(QFont("Arial", 14))
        output_label.setAlignment(Qt.AlignLeft)

        # Copy button
        self.copy_button = QPushButton("Copy Output")
        self.copy_button.setFont(QFont("Arial", 12))
        self.copy_button.clicked.connect(self.copy_output)

        # Add widgets to layout
        main_layout.addLayout(button_layout)
        main_layout.addLayout(slider_layout)
        main_layout.addWidget(input_label)
        main_layout.addWidget(self.text_display)
        main_layout.addWidget(output_label)
        main_layout.addWidget(self.output_display)
        main_layout.addWidget(self.copy_button, alignment=Qt.AlignCenter)

        # Set central widget
        container = QWidget()
        container.setLayout(main_layout)
        self.setCentralWidget(container)

    def update_summary_length(self):
        self.max_summary_length = self.length_slider.value()
        self.slider_value_label.setText(f"{self.max_summary_length} words")

    def open_file(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getOpenFileName(self, "Open File", "", "PDF Files (*.pdf);;Word Files (*.docx);;All Files (*)", options=options)

        if file_path:
            self.text_display.clear()
            self.output_display.clear()

            if file_path.endswith(".pdf"):
                self.pdf_text = self.extract_text_from_pdf(file_path)
            elif file_path.endswith(".docx"):
                self.pdf_text = self.extract_text_from_docx(file_path)
            else:
                self.show_error("Unsupported file format.")
                return

            self.text_display.setText(self.pdf_text)
            self.summarize_button.setEnabled(True)
            self.extract_stats_button.setEnabled(True)

    def extract_text_from_pdf(self, pdf_path):
        doc = fitz.open(pdf_path)
        text_chunks = []
        for page in doc:
            text = page.get_text()
            if not text.strip():
                # Perform OCR if no text is found
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text = pytesseract.image_to_string(img)
            text_chunks.append(text)
        return "\n".join(text_chunks)

    def extract_text_from_docx(self, docx_path):
        doc = Document(docx_path)
        text_chunks = []
        for paragraph in doc.paragraphs:
            text_chunks.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_chunks.append(cell.text)

        return "\n".join(text_chunks)

    def summarize_text(self):
        if hasattr(self, 'pdf_text') and self.pdf_text:
            if not self.summarizer:
                self.show_error("Summarization is unavailable. Install PyTorch or TensorFlow.")
                return

            chunks = self.split_text_into_chunks(self.pdf_text, 700)  # Reduce chunk size for efficiency
            summaries = []
            for chunk in chunks:
                summary = self.summarizer(chunk, max_length=self.max_summary_length, min_length=20, do_sample=False)[0]['summary_text']
                summaries.append(summary)
            self.output_display.setText("\n\n".join(summaries))

    def split_text_into_chunks(self, text, max_chunk_size):
        words = text.split()
        chunks = []
        current_chunk = []
        current_length = 0

        for word in words:
            if current_length + len(word) + 1 > max_chunk_size:
                chunks.append(" ".join(current_chunk))
                current_chunk = []
                current_length = 0
            current_chunk.append(word)
            current_length += len(word) + 1

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks

    def extract_stats(self):
        if hasattr(self, 'pdf_text') and self.pdf_text:
            if not self.nlp:
                self.show_error("Keyword extraction is unavailable due to missing spaCy model.")
                return

            # Example: Extract numbers and keywords
            numbers = re.findall(r'\b\d+\b', self.pdf_text)
            doc = self.nlp(self.pdf_text)
            keywords = [token.text for token in doc if token.is_alpha and not token.is_stop]

            stats = f"Numbers Found: {', '.join(numbers)}\n\nKeywords: {', '.join(keywords[:50])}"
            self.output_display.setText(stats)

    def copy_output(self):
        output_text = self.output_display.toPlainText()
        QApplication.clipboard().setText(output_text)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = PDFSummarizerApp()
    window.show()
    sys.exit(app.exec_())
